from docx import Document
from io import BytesIO
def build_docx(dataset: dict):
    doc = Document()
    doc.add_heading(dataset.get("title","Relatório Fiscal"), 0)
    doc.add_heading("KPIs", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Métrica"
    table.rows[0].cells[1].text = "Valor"
    for k in dataset.get("kpis", []):
        row = table.add_row().cells
        row[0].text = str(k.get("label",""))
        row[1].text = str(k.get("value",""))
    doc.add_heading("Inconsistências", level=1)
    table2 = doc.add_table(rows=1, cols=4)
    table2.rows[0].cells[0].text = "Campo"
    table2.rows[0].cells[1].text = "Código"
    table2.rows[0].cells[2].text = "Severidade"
    table2.rows[0].cells[3].text = "Descrição"
    for i in dataset.get("compliance",{}).get("inconsistencies",[]):
        row = table2.add_row().cells
        row[0].text = str(i.get("field",""))
        row[1].text = str(i.get("code",""))
        row[2].text = str(i.get("severity",""))
        row[3].text = str(i.get("message",""))
    bio = BytesIO()
    doc.save(bio); bio.seek(0)
    return bio.getvalue(), "relatorio.docx"
